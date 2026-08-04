#!/usr/bin/env python3
"""Publish ServicePro Markdown as enterprise Markdown, DOCX, and PDF artifacts.

The source file is never modified. This first version intentionally supports the
Markdown constructs used by the ServicePro documentation corpus and is designed
to be extended batch-by-batch as the library is published.
"""

from __future__ import annotations

import argparse
import html
import re
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    KeepTogether,
    ListFlowable,
    ListItem,
    PageBreak,
    Paragraph,
    Preformatted,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
    Image as RLImage,
)
from PIL import Image as PILImage, ImageDraw, ImageFont

NAVY = "17365D"
BLUE = "2E74B5"
SKY = "DCE6F1"
PALE = "F3F6FA"
INK = "243447"
MUTED = "66788A"
WHITE = "FFFFFF"
GOLD = "C7922B"
TOTAL_DXA = 9360


def add_bookmark(paragraph, name: str, bookmark_id: int):
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    paragraph_properties = paragraph._p.find(qn("w:pPr"))
    insert_at = 1 if paragraph_properties is not None else 0
    paragraph._p.insert(insert_at, start)
    paragraph._p.append(end)


def add_internal_hyperlink(paragraph, label: str, anchor: str):
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), anchor)
    run = OxmlElement("w:r")
    props = OxmlElement("w:rPr")
    style = OxmlElement("w:rStyle")
    style.set(qn("w:val"), "Hyperlink")
    props.append(style)
    run.append(props)
    text = OxmlElement("w:t")
    text.text = label
    run.append(text)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


class NavigableDocTemplate(SimpleDocTemplate):
    def afterFlowable(self, flowable):
        anchor = getattr(flowable, "_bookmark_name", None)
        if not anchor:
            return
        title = getattr(flowable, "_bookmark_title", anchor)
        requested_level = getattr(flowable, "_outline_level", 0)
        previous_level = getattr(self, "_last_outline_level", -1)
        level = max(0, min(requested_level, previous_level + 1))
        self.canv.bookmarkPage(anchor)
        self.canv.addOutlineEntry(title, anchor, level=level, closed=False)
        self._last_outline_level = level


def clean_inline(text: str) -> str:
    text = re.sub(r"\[([^]]+)]\([^)]+\)", r"\1", text)
    text = text.replace("**", "").replace("__", "").replace("`", "")
    return text.strip()


def split_blocks(text: str):
    lines = text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]
        if line.startswith("```"):
            language = line[3:].strip()
            buf = []
            i += 1
            while i < len(lines) and not lines[i].startswith("```"):
                buf.append(lines[i])
                i += 1
            yield ("code", language, "\n".join(buf))
        elif re.match(r"^#{1,6}\s+", line):
            m = re.match(r"^(#{1,6})\s+(.*)$", line)
            yield ("heading", len(m.group(1)), clean_inline(m.group(2)))
        elif line.strip() == "---":
            yield ("rule",)
        elif line.startswith("|") and i + 1 < len(lines) and re.match(r"^\|?\s*:?-+", lines[i + 1]):
            rows = []
            while i < len(lines) and lines[i].startswith("|"):
                cells = [clean_inline(c) for c in lines[i].strip().strip("|").split("|")]
                if not all(re.fullmatch(r":?-+:?", c.replace(" ", "")) for c in cells):
                    rows.append(cells)
                i += 1
            yield ("table", rows)
            continue
        elif re.match(r"^\s*[-*]\s+", line):
            items = []
            while i < len(lines) and re.match(r"^\s*[-*]\s+", lines[i]):
                items.append(clean_inline(re.sub(r"^\s*[-*]\s+", "", lines[i])))
                i += 1
            yield ("bullets", items)
            continue
        elif re.match(r"^\s*\d+\.\s+", line):
            items = []
            while i < len(lines) and re.match(r"^\s*\d+\.\s+", lines[i]):
                items.append(clean_inline(re.sub(r"^\s*\d+\.\s+", "", lines[i])))
                i += 1
            yield ("numbers", items)
            continue
        elif not line.strip():
            yield ("space",)
        else:
            buf = [line.strip()]
            while i + 1 < len(lines):
                nxt = lines[i + 1]
                if not nxt.strip() or nxt.startswith(("#", "|", "```", "---")) or re.match(r"^\s*(?:[-*]|\d+\.)\s+", nxt):
                    break
                i += 1
                buf.append(nxt.strip())
            yield ("paragraph", " ".join(buf))
        i += 1


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tag = OxmlElement("w:tblHeader")
    tag.set(qn("w:val"), "true")
    tr_pr.append(tag)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths)))
    ind = tbl_pr.find(qn("w:tblInd"))
    if ind is None:
        ind = OxmlElement("w:tblInd")
        tbl_pr.append(ind)
    ind.set(qn("w:type"), "dxa")
    ind.set(qn("w:w"), "120")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(widths[idx]))
            cell.width = Inches(widths[idx] / 1440)
            set_cell_margins(cell)


def add_field(paragraph, instruction):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    text = OxmlElement("w:instrText")
    text.set(qn("xml:space"), "preserve")
    text.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, text, separate, end])


def style_doc(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.82)
    section.bottom_margin = Inches(0.78)
    section.left_margin = Inches(0.82)
    section.right_margin = Inches(0.82)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)
    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.12
    for name, size, color, before, after in (
        ("Heading 1", 18, NAVY, 18, 7),
        ("Heading 2", 14, BLUE, 14, 5),
        ("Heading 3", 11.5, NAVY, 10, 4),
    ):
        st = doc.styles[name]
        st.font.name = "Aptos Display"
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor.from_string(color)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True
    for name in ("List Bullet", "List Number"):
        st = doc.styles[name]
        st.font.name = "Aptos"
        st.font.size = Pt(10.5)
        st.paragraph_format.left_indent = Inches(0.38)
        st.paragraph_format.first_line_indent = Inches(-0.19)
        st.paragraph_format.space_after = Pt(3)


def add_rich_text(paragraph, text):
    pattern = re.compile(r"(\*\*.+?\*\*|`.+?`|\[[^]]+\]\([^)]+\))")
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            paragraph.add_run(text[pos:match.start()])
        token = match.group(0)
        if token.startswith("**"):
            r = paragraph.add_run(token[2:-2]); r.bold = True
        elif token.startswith("`"):
            r = paragraph.add_run(token[1:-1]); r.font.name = "Consolas"; r.font.size = Pt(9)
        else:
            label = re.match(r"\[([^]]+)]", token).group(1)
            r = paragraph.add_run(label); r.font.color.rgb = RGBColor.from_string(BLUE); r.underline = True
        pos = match.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def build_docx(source_text: str, output: Path, title: str, subtitle: str, document_type: str = "ServicePro Documentation"):
    blocks = list(split_blocks(source_text))
    navigation = []
    skipped = 0
    for block in blocks:
        if block[0] == "heading" and block[1] <= 2 and skipped < 2:
            skipped += 1
            continue
        if block[0] == "heading":
            level = min(3, max(1, block[1] - 1))
            navigation.append((level, block[2], f"section_{len(navigation) + 1:03d}"))
    doc = Document()
    style_doc(doc)
    section = doc.sections[0]
    header = section.header.paragraphs[0]
    header.text = f"SERVICEPRO  |  {document_type.upper()}"
    header.style = doc.styles["Caption"]
    header.runs[0].font.color.rgb = RGBColor.from_string(MUTED)
    header.runs[0].font.size = Pt(8)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = footer.add_run("ServicePro Documentation Library  |  ")
    r.font.size = Pt(8); r.font.color.rgb = RGBColor.from_string(MUTED)
    add_field(footer, "PAGE")

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(115)
    p.paragraph_format.space_after = Pt(10)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(document_type.upper())
    r.bold = True; r.font.name = "Aptos"; r.font.size = Pt(10); r.font.color.rgb = RGBColor.from_string(GOLD)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run(title)
    r.bold = True; r.font.name = "Aptos Display"; r.font.size = Pt(30); r.font.color.rgb = RGBColor.from_string(NAVY)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(80)
    r = p.add_run(subtitle)
    r.font.name = "Aptos"; r.font.size = Pt(14); r.font.color.rgb = RGBColor.from_string(BLUE)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"Publication edition  |  {date.today().isoformat()}")
    r.font.size = Pt(9.5); r.font.color.rgb = RGBColor.from_string(MUTED)
    doc.add_page_break()

    document_control_heading = doc.add_heading("Document Control", level=1)
    add_bookmark(document_control_heading, "document_control", 1)
    table = doc.add_table(rows=4, cols=2)
    meta = [("Purpose", "Enterprise platform overview and buyer evaluation reference"),
            ("Audience", "Business leaders, platform administrators, evaluators, partners, and technical stakeholders"),
            ("Scope", "Capabilities, architecture, security, deployment, adoption, outcomes, and terminology"),
            ("Source", "ServicePro repository documentation; technical meaning preserved")]
    for row, values in zip(table.rows, meta):
        for idx, value in enumerate(values):
            row.cells[idx].text = value
            row.cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            shade(row.cells[idx], SKY if idx == 0 else WHITE)
            if idx == 0:
                row.cells[idx].paragraphs[0].runs[0].bold = True
    set_table_geometry(table, [2100, 7260])
    doc.add_paragraph("Use the linked table of contents below or Word's Navigation Pane to move through this document.")
    toc_heading = doc.add_heading("Table of Contents", level=1)
    add_bookmark(toc_heading, "table_of_contents", 2)
    document_control_entry = doc.add_paragraph()
    document_control_entry.paragraph_format.space_after = Pt(3)
    add_internal_hyperlink(document_control_entry, "Document Control", "document_control")
    for level, heading_title, anchor in navigation:
        entry = doc.add_paragraph()
        entry.paragraph_format.left_indent = Inches(0.22 * (level - 1))
        entry.paragraph_format.space_after = Pt(3)
        add_internal_hyperlink(entry, heading_title, anchor)
    doc.add_page_break()

    skip_first_title = 0
    heading_cursor = 0
    for block in blocks:
        kind = block[0]
        if kind == "heading" and block[1] <= 2 and skip_first_title < 2:
            skip_first_title += 1
            continue
        if kind == "heading":
            level = min(3, max(1, block[1] - 1))
            heading = doc.add_heading(block[2], level=level)
            add_bookmark(heading, navigation[heading_cursor][2], heading_cursor + 3)
            heading_cursor += 1
        elif kind == "paragraph":
            add_rich_text(doc.add_paragraph(), block[1])
        elif kind in ("bullets", "numbers"):
            style = "List Bullet" if kind == "bullets" else "List Number"
            for item in block[1]:
                add_rich_text(doc.add_paragraph(style=style), item)
        elif kind == "table" and block[1]:
            rows = block[1]
            cols = max(len(r) for r in rows)
            table = doc.add_table(rows=len(rows), cols=cols)
            for ridx, values in enumerate(rows):
                for cidx in range(cols):
                    cell = table.cell(ridx, cidx)
                    cell.text = values[cidx] if cidx < len(values) else ""
                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    if ridx == 0:
                        shade(cell, NAVY)
                        for run in cell.paragraphs[0].runs:
                            run.bold = True; run.font.color.rgb = RGBColor.from_string(WHITE)
                    elif ridx % 2 == 0:
                        shade(cell, PALE)
                    for run in cell.paragraphs[0].runs:
                        run.font.size = Pt(8.2 if cols >= 4 else 9)
            set_repeat_table_header(table.rows[0])
            weights = [max(8, max(len(r[c]) if c < len(r) else 0 for r in rows)) for c in range(cols)]
            widths = [max(900, round(TOTAL_DXA * w / sum(weights))) for w in weights]
            widths[-1] += TOTAL_DXA - sum(widths)
            set_table_geometry(table, widths)
            doc.add_paragraph().paragraph_format.space_after = Pt(2)
        elif kind == "code" and block[1].lower() == "mermaid" and (output.parent / "servicepro-platform-architecture.png").exists():
            diagram = output.parent / "servicepro-platform-architecture.png"
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.keep_with_next = True
            p.add_run().add_picture(str(diagram), width=Inches(6.35))
            c = doc.add_paragraph("Figure 1. ServicePro platform architecture and request flow")
            c.alignment = WD_ALIGN_PARAGRAPH.CENTER
            c.style = doc.styles["Caption"]
        elif kind == "code":
            p = doc.add_paragraph()
            p.style = doc.styles["No Spacing"]
            p.paragraph_format.left_indent = Inches(0.18)
            p.paragraph_format.space_before = Pt(5); p.paragraph_format.space_after = Pt(7)
            r = p.add_run(block[2]); r.font.name = "Consolas"; r.font.size = Pt(7.5); r.font.color.rgb = RGBColor.from_string(INK)
            p_pr = p._p.get_or_add_pPr(); shd = OxmlElement("w:shd"); shd.set(qn("w:fill"), PALE); p_pr.append(shd)
        elif kind == "rule":
            p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(3)
    doc.core_properties.title = title
    doc.core_properties.subject = subtitle
    doc.core_properties.author = "ServicePro"
    doc.core_properties.keywords = "ServicePro, enterprise platform, field service, SaaS"
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def pdf_text(text):
    return html.escape(clean_inline(text)).replace("\n", "<br/>")


def build_pdf(source_text: str, output: Path, title: str, subtitle: str, document_type: str = "ServicePro Documentation"):
    blocks = list(split_blocks(source_text))
    navigation = []
    skipped = 0
    for block in blocks:
        if block[0] == "heading" and block[1] <= 2 and skipped < 2:
            skipped += 1
            continue
        if block[0] == "heading":
            level = min(3, max(1, block[1] - 1))
            navigation.append((level, block[2], f"pdf_section_{len(navigation) + 1:03d}"))
    styles = getSampleStyleSheet()
    body = ParagraphStyle("Body", parent=styles["BodyText"], fontName="Helvetica", fontSize=9.2, leading=12.2, textColor=colors.HexColor("#243447"), spaceAfter=6)
    h1 = ParagraphStyle("H1", parent=styles["Heading1"], fontName="Helvetica-Bold", fontSize=17, leading=21, textColor=colors.HexColor("#17365D"), spaceBefore=14, spaceAfter=7, keepWithNext=True)
    h2 = ParagraphStyle("H2", parent=styles["Heading2"], fontName="Helvetica-Bold", fontSize=13, leading=16, textColor=colors.HexColor("#2E74B5"), spaceBefore=11, spaceAfter=5, keepWithNext=True)
    h3 = ParagraphStyle("H3", parent=styles["Heading3"], fontName="Helvetica-Bold", fontSize=10.5, leading=13, textColor=colors.HexColor("#17365D"), spaceBefore=8, spaceAfter=4, keepWithNext=True)
    code = ParagraphStyle("Code", parent=styles["Code"], fontName="Courier", fontSize=6.3, leading=8, backColor=colors.HexColor("#F3F6FA"), borderPadding=6, spaceAfter=7)

    def page(canvas, doc):
        canvas.saveState(); canvas.setStrokeColor(colors.HexColor("#DCE6F1")); canvas.line(0.82*inch, 10.43*inch, 7.68*inch, 10.43*inch)
        canvas.setFont("Helvetica", 7.5); canvas.setFillColor(colors.HexColor("#66788A")); canvas.drawString(0.82*inch, 10.56*inch, f"SERVICEPRO  |  {document_type.upper()}")
        canvas.drawRightString(7.68*inch, 0.42*inch, f"ServicePro Documentation Library  |  {doc.page}"); canvas.restoreState()

    output.parent.mkdir(parents=True, exist_ok=True)
    doc = NavigableDocTemplate(str(output), pagesize=LETTER, rightMargin=.82*inch, leftMargin=.82*inch, topMargin=.75*inch, bottomMargin=.68*inch, title=title, author="ServicePro")
    document_control = Paragraph('<a name="pdf_document_control"/>Document Control', h1)
    document_control._bookmark_name = "pdf_document_control"
    document_control._bookmark_title = "Document Control"
    document_control._outline_level = 0
    story = [Spacer(1, 1.42*inch), Paragraph(pdf_text(document_type.upper()), ParagraphStyle("K", parent=body, alignment=TA_CENTER, textColor=colors.HexColor("#C7922B"), fontName="Helvetica-Bold", fontSize=9)), Spacer(1, 10), Paragraph(pdf_text(title), ParagraphStyle("T", parent=h1, alignment=TA_CENTER, fontSize=28, leading=32, spaceAfter=12)), Paragraph(pdf_text(subtitle), ParagraphStyle("S", parent=body, alignment=TA_CENTER, fontSize=13, leading=17, textColor=colors.HexColor("#2E74B5"))), Spacer(1, 1.1*inch), Paragraph(f"Publication edition  |  {date.today().isoformat()}", ParagraphStyle("D", parent=body, alignment=TA_CENTER, textColor=colors.HexColor("#66788A"))), PageBreak(), document_control]
    meta = [[Paragraph("Purpose", h3), Paragraph("Enterprise platform overview and buyer evaluation reference", body)], [Paragraph("Audience", h3), Paragraph("Business leaders, platform administrators, evaluators, partners, and technical stakeholders", body)], [Paragraph("Scope", h3), Paragraph("Capabilities, architecture, security, deployment, adoption, outcomes, and terminology", body)], [Paragraph("Source", h3), Paragraph("ServicePro repository documentation; technical meaning preserved", body)]]
    t = Table(meta, colWidths=[1.35*inch, 5.15*inch], repeatRows=0)
    t.setStyle(TableStyle([("BACKGROUND",(0,0),(0,-1),colors.HexColor("#DCE6F1")),("GRID",(0,0),(-1,-1),.35,colors.HexColor("#AAB7C4")),("VALIGN",(0,0),(-1,-1),"MIDDLE"),("LEFTPADDING",(0,0),(-1,-1),7),("RIGHTPADDING",(0,0),(-1,-1),7),("TOPPADDING",(0,0),(-1,-1),5),("BOTTOMPADDING",(0,0),(-1,-1),5)])); story += [t, Spacer(1, 12), Paragraph("Table of Contents", h1), Paragraph('<a href="#pdf_document_control" color="#2E74B5">Document Control</a>', body)]
    for level, heading_title, anchor in navigation:
        toc_style = ParagraphStyle(f"TOC{level}", parent=body, leftIndent=(level - 1) * 14, textColor=colors.HexColor("#2E74B5"), spaceAfter=4)
        story.append(Paragraph(f'<a href="#{anchor}" color="#2E74B5">{html.escape(heading_title)}</a>', toc_style))
    story.append(PageBreak())
    skip = 0
    heading_cursor = 0
    for b in blocks:
        if b[0] == "heading" and b[1] <= 2 and skip < 2: skip += 1; continue
        if b[0] == "heading":
            level, heading_title, anchor = navigation[heading_cursor]
            heading = Paragraph(f'<a name="{anchor}"/>{pdf_text(b[2])}', [h1,h2,h3][level - 1])
            heading._bookmark_name = anchor
            heading._bookmark_title = heading_title
            heading._outline_level = level - 1
            story.append(heading)
            heading_cursor += 1
        elif b[0] == "paragraph": story.append(Paragraph(pdf_text(b[1]), body))
        elif b[0] in ("bullets","numbers"):
            items=[ListItem(Paragraph(pdf_text(x),body), leftIndent=10) for x in b[1]]
            story.append(ListFlowable(items, bulletType="bullet" if b[0]=="bullets" else "1", leftIndent=20, bulletFontSize=7, spaceAfter=5))
        elif b[0] == "table" and b[1]:
            rows=[[Paragraph(pdf_text(c), ParagraphStyle("Cell",parent=body,fontSize=6.7,leading=8.2,textColor=colors.white if ri==0 else colors.HexColor("#243447"),spaceAfter=0)) for c in row] for ri,row in enumerate(b[1])]
            cols=max(len(r) for r in rows); rows=[r+[Paragraph("",body)]*(cols-len(r)) for r in rows]
            weights=[max(8,max(len(clean_inline(b[1][r][c])) if c<len(b[1][r]) else 0 for r in range(len(b[1]))) ) for c in range(cols)]
            widths=[6.5*inch*w/sum(weights) for w in weights]
            tab=Table(rows,colWidths=widths,repeatRows=1,hAlign="LEFT")
            tab.setStyle(TableStyle([("BACKGROUND",(0,0),(-1,0),colors.HexColor("#17365D")),("ROWBACKGROUNDS",(0,1),(-1,-1),[colors.white,colors.HexColor("#F3F6FA")]),("GRID",(0,0),(-1,-1),.3,colors.HexColor("#AAB7C4")),("VALIGN",(0,0),(-1,-1),"MIDDLE"),("LEFTPADDING",(0,0),(-1,-1),4),("RIGHTPADDING",(0,0),(-1,-1),4),("TOPPADDING",(0,0),(-1,-1),4),("BOTTOMPADDING",(0,0),(-1,-1),4)])); story += [tab,Spacer(1,6)]
        elif b[0] == "code" and b[1].lower() == "mermaid" and (output.parent / "servicepro-platform-architecture.png").exists():
            diagram = output.parent / "servicepro-platform-architecture.png"
            story += [RLImage(str(diagram), width=6.35*inch, height=4.64*inch), Paragraph("Figure 1. ServicePro platform architecture and request flow", ParagraphStyle("Fig",parent=body,alignment=TA_CENTER,fontSize=8,textColor=colors.HexColor("#66788A"),spaceAfter=8))]
        elif b[0] == "code": story.append(Preformatted(b[2],code,maxLineLength=105))
        elif b[0] == "space": story.append(Spacer(1,2))
    doc.build(story,onFirstPage=page,onLaterPages=page)


def polished_markdown(source: str, title: str, subtitle: str, document_type: str = "ServicePro Documentation") -> str:
    lines = source.splitlines()
    body = "\n".join(line.rstrip() for line in lines[2:]).lstrip()
    front = f"""---
title: "{title}"
subtitle: "{subtitle}"
document_type: "{document_type}"
audience:
  - Business leaders
  - Platform administrators
  - Buyers and evaluators
  - Partners and technical stakeholders
status: "Publication edition"
published: "{date.today().isoformat()}"
source_of_truth: "ServicePro repository"
---

# {title}

> **{document_type}**
> {subtitle}

## Document Control

| Field | Detail |
|---|---|
| Purpose | Enterprise platform overview and buyer evaluation reference |
| Audience | Business leaders, platform administrators, evaluators, partners, and technical stakeholders |
| Scope | Capabilities, architecture, security, deployment, adoption, outcomes, and terminology |
| Source | ServicePro repository documentation; technical meaning preserved |

> [!NOTE]
> This publication edition improves navigation, document metadata, and cross-format consistency. Product and technical claims remain those of the source document.

"""
    if body.strip():
        return front.rstrip() + "\n\n" + body.strip() + "\n"
    return front.rstrip() + "\n"


def build_architecture_diagram(output: Path):
    """Create a crisp, repo-independent architecture figure from source semantics."""
    width, height = 1800, 1320
    image = PILImage.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    try:
        regular = ImageFont.truetype("arial.ttf", 27)
        small = ImageFont.truetype("arial.ttf", 23)
        bold = ImageFont.truetype("arialbd.ttf", 30)
        title_font = ImageFont.truetype("arialbd.ttf", 42)
    except OSError:
        regular = small = bold = title_font = ImageFont.load_default()

    draw.text((90, 55), "ServicePro Platform Architecture", fill="#17365D", font=title_font)
    draw.text((90, 110), "Multi-application request flow with shared security, services, and data", fill="#66788A", font=regular)

    layers = [
        ("CLIENT APPLICATIONS", 185, ["Tenant App", "Customer Portal", "Admin App", "Mobile App"], "#DCE6F1"),
        ("API & SECURITY", 440, ["REST API", "JWT Auth", "CORS", "Rate Limiting", "Tenant Isolation", "Audit"], "#E8F1F8"),
        ("BUSINESS SERVICES", 695, ["Jobs", "Scheduling", "CRM", "Inventory", "Billing", "Knowledge", "Notifications", "Workflows", "AI", "Marketing", "Website", "Marketplace"], "#F3F6FA"),
        ("DATA LAYER", 1050, ["PostgreSQL (production)", "JSON Store (development)"], "#FFF5DF"),
    ]
    for idx, (label, y, nodes, fill) in enumerate(layers):
        draw.rounded_rectangle((75, y, 1725, y + (235 if idx == 2 else 165)), radius=18, fill=fill, outline="#AAB7C4", width=3)
        draw.text((100, y + 18), label, fill="#2E74B5", font=bold)
        cols = 4 if idx in (0, 2) else (3 if idx == 1 else 2)
        box_w = 350 if cols == 4 else (460 if cols == 3 else 670)
        gap = (1540 - cols * box_w) / max(1, cols - 1)
        for n, node in enumerate(nodes):
            row, col = divmod(n, cols)
            x = 125 + col * (box_w + gap)
            by = y + 72 + row * 75
            draw.rounded_rectangle((x, by, x + box_w, by + 52), radius=10, fill="white", outline="#7F95AA", width=2)
            bbox = draw.textbbox((0, 0), node, font=small)
            draw.text((x + (box_w - (bbox[2]-bbox[0]))/2, by + 12), node, fill="#243447", font=small)
        if idx < len(layers) - 1:
            next_y = layers[idx + 1][1]
            cx = width // 2
            draw.line((cx, y + (235 if idx == 2 else 165), cx, next_y - 12), fill="#2E74B5", width=6)
            draw.polygon([(cx - 12, next_y - 28), (cx + 12, next_y - 28), (cx, next_y - 8)], fill="#2E74B5")
    draw.text((90, 1270), "Security and tenant isolation are enforced before business-service access.", fill="#66788A", font=regular)
    output.parent.mkdir(parents=True, exist_ok=True)
    image.save(output, "PNG")


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("source", type=Path)
    ap.add_argument("output_dir", type=Path)
    args = ap.parse_args()
    source = args.source.read_text(encoding="utf-8")
    title = "ServicePro Enterprise Platform"
    subtitle = "Unified Field Service, Business Operations, Customer Experience, and Digital Growth"
    stem = "servicepro-enterprise-platform-overview"
    args.output_dir.mkdir(parents=True, exist_ok=True)
    md = args.output_dir / f"{stem}.md"
    docx = args.output_dir / f"{stem}.docx"
    pdf = args.output_dir / f"{stem}.pdf"
    build_architecture_diagram(args.output_dir / "servicepro-platform-architecture.png")
    md.write_text(polished_markdown(source,title,subtitle),encoding="utf-8")
    build_docx(source,docx,title,subtitle)
    build_pdf(source,pdf,title,subtitle)
    print(md); print(docx); print(pdf)


if __name__ == "__main__":
    main()
