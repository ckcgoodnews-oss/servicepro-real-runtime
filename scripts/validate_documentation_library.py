#!/usr/bin/env python3
"""Validate every deployable documentation artifact and manifest link."""

import json
import sys
import zipfile
from pathlib import Path
from xml.etree import ElementTree

from docx import Document
from pypdf import PdfReader


def main():
    repo = Path(sys.argv[1] if len(sys.argv) > 1 else ".").resolve()
    public = repo / "apps/web/public"
    index_path = public / "documentation/library/index.json"
    index = json.loads(index_path.read_text(encoding="utf-8"))
    documents = index["documents"]
    assert index["documentCount"] == len(documents) == 639
    seen = set()
    page_count = 0
    for number, record in enumerate(documents, 1):
        assert record["id"] not in seen, f"duplicate id: {record['id']}"
        seen.add(record["id"])
        paths = {}
        for kind in ("markdown", "docx", "pdf"):
            path = public / record[f"{kind}Url"].lstrip("/")
            assert path.is_file() and path.stat().st_size > 0, f"missing {kind}: {path}"
            paths[kind] = path
        markdown = paths["markdown"].read_text(encoding="utf-8")
        assert markdown.startswith("---\n") and f"# {record['title']}" in markdown
        word = Document(paths["docx"])
        assert len(word.paragraphs) > 0 and any(p.text.strip() for p in word.paragraphs)
        with zipfile.ZipFile(paths["docx"]) as package:
            xml = ElementTree.fromstring(package.read("word/document.xml"))
            namespace = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
            bookmarks = xml.findall(".//w:bookmarkStart", namespace)
            hyperlinks = [node for node in xml.findall(".//w:hyperlink", namespace) if node.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}anchor")]
            assert bookmarks and hyperlinks, f"missing Word TOC navigation: {paths['docx']}"
            bookmark_names = {node.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}name") for node in bookmarks}
            hyperlink_anchors = {node.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}anchor") for node in hyperlinks}
            assert hyperlink_anchors <= bookmark_names, f"Word TOC links missing bookmark targets: {paths['docx']}"
            for paragraph in xml.findall(".//w:p", namespace):
                children = list(paragraph)
                properties_positions = [idx for idx, child in enumerate(children) if child.tag == "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pPr"]
                assert not properties_positions or properties_positions == [0], f"invalid Word paragraph property order: {paths['docx']}"
        pdf = PdfReader(paths["pdf"])
        assert len(pdf.pages) >= 1
        assert pdf.outline, f"missing PDF outline: {paths['pdf']}"
        assert any(page.get("/Annots") for page in pdf.pages), f"missing PDF TOC links: {paths['pdf']}"
        page_count += len(pdf.pages)
        if number % 100 == 0:
            print(f"validated {number}/{len(documents)}")
    expected = {".md": 639, ".docx": 639, ".pdf": 639}
    files_root = public / "documentation/library/files"
    for suffix, count in expected.items():
        actual = len(list(files_root.rglob(f"*{suffix}")))
        assert actual == count, f"expected {count} {suffix} files, found {actual}"
    print(f"validated {len(documents)} documents, {len(documents) * 3} formats, {page_count} PDF pages")


if __name__ == "__main__":
    main()
